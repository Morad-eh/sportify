# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUTPUT_DIR = r'C:\Users\morad\sportify\documents_tfe'
os.makedirs(OUTPUT_DIR, exist_ok=True)

VERT = RGBColor(0x1A, 0x8C, 0x4E)
NOIR = RGBColor(0x0D, 0x2E, 0x1A)


def set_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = VERT
        run.font.size = Pt(15) if level == 1 else Pt(12)
    return p


def add_para(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def title_page(doc, title, subtitle):
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('SPORTIFY')
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = VERT

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(title)
    r2.bold = True
    r2.font.size = Pt(18)
    r2.font.color.rgb = NOIR

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(subtitle)
    r3.font.size = Pt(12)

    for _ in range(3):
        doc.add_paragraph()

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run("EL HALIMI Morad — Bac Info Appli\nAnnée académique 2025-2026")
    r4.font.size = Pt(11)

    doc.add_page_break()


# ============================================================
#  T16 - STRATEGIE DE SECURITE
# ============================================================
def create_t16():
    doc = Document()
    title_page(doc, "Stratégie de Sécurité", "Travail 16 — Application Web Sportify")

    set_heading(doc, "1. Introduction")
    add_para(doc,
        "Dans le cadre de mon projet TFE, j'ai développé une application web appelée Sportify "
        "qui permet de réserver des terrains sportifs indoor. Comme l'application gère des "
        "comptes utilisateurs et des paiements, la sécurité est très importante. "
        "Dans ce document, j'explique les mesures que j'ai mises en place pour protéger "
        "l'application et les données des utilisateurs.")

    set_heading(doc, "2. Connexion et gestion des accès")
    add_para(doc,
        "Pour que les utilisateurs puissent se connecter à Sportify, j'ai utilisé le système "
        "d'authentification intégré de Django. Ce système est reconnu comme sécurisé et "
        "gère automatiquement plusieurs choses importantes :")
    add_bullet(doc, "Les mots de passe ne sont jamais stockés en clair dans la base de données. Django les transforme avec un algorithme de hachage (PBKDF2-SHA256) ce qui fait que même si quelqu'un accède à la base de données, il ne peut pas lire les mots de passe.")
    add_bullet(doc, "Chaque utilisateur a une session unique et sécurisée. Quand il se déconnecte, la session est immédiatement supprimée.")
    add_bullet(doc, "Django protège automatiquement les formulaires contre les attaques CSRF (une attaque où un site malveillant envoie des requêtes à la place de l'utilisateur).")

    add_para(doc,
        "Dans Sportify, j'ai créé 3 rôles différents : Membre, Propriétaire de terrain et "
        "Administrateur. Chaque rôle n'a accès qu'aux fonctionnalités dont il a besoin. "
        "Par exemple, un membre ne peut pas accéder à l'interface d'administration, "
        "et un utilisateur non connecté ne peut pas faire de réservation.")

    set_heading(doc, "3. Les 10 risques OWASP et comment je les gère")
    add_para(doc,
        "L'OWASP TOP 10 est une liste publiée par des experts en sécurité web qui recense "
        "les 10 types d'attaques les plus fréquentes sur les sites internet. "
        "Voici comment Sportify se protège contre chacune :")

    owasp = [
        ("A01 — Mauvais contrôle des accès",
         "Dans Sportify, chaque page sensible est protégée. Si quelqu'un essaie d'accéder "
         "au tableau de bord sans être connecté, il est automatiquement redirigé vers la page "
         "de connexion. Un membre ne peut pas voir les réservations d'un autre utilisateur."),
        ("A02 — Mauvaise protection des données",
         "Les mots de passe sont hachés (transformés de façon irréversible). "
         "Le site fonctionne en HTTPS sur PythonAnywhere, ce qui chiffre toutes les "
         "communications. Les clés secrètes de l'API Stripe ne sont jamais visibles côté utilisateur."),
        ("A03 — Injections SQL",
         "Django utilise un système appelé ORM pour communiquer avec la base de données. "
         "Concrètement, cela veut dire que les données saisies par les utilisateurs ne sont "
         "jamais directement insérées dans les requêtes SQL. Le risque d'injection SQL est "
         "donc éliminé automatiquement."),
        ("A04 — Conception non sécurisée",
         "Dès le début du projet, j'ai réfléchi à la sécurité avant de coder. "
         "Par exemple, j'ai défini les rôles utilisateurs et les droits d'accès "
         "dans la conception des modèles, avant même de créer les pages."),
        ("A05 — Mauvaise configuration du serveur",
         "En production sur PythonAnywhere, j'ai désactivé le mode DEBUG de Django "
         "(qui afficherait des infos sensibles en cas d'erreur). J'ai aussi restreint "
         "ALLOWED_HOSTS à mon domaine uniquement."),
        ("A06 — Librairies obsolètes",
         "J'utilise Django 5.0.14, une version récente et maintenue activement. "
         "Toutes les dépendances sont listées dans le fichier requirements.txt avec "
         "des versions précises pour éviter les mises à jour non contrôlées."),
        ("A07 — Mauvaise gestion des connexions",
         "Django gère les sessions de manière sécurisée. Un utilisateur qui ne se connecte "
         "pas avec le bon mot de passe ne peut pas accéder au site. "
         "La déconnexion invalide immédiatement la session côté serveur."),
        ("A08 — Intégrité du code",
         "Le code source de Sportify est versionné sur GitHub. Cela permet de savoir "
         "exactement ce qui a été modifié et quand. Toutes les dépendances sont fixées "
         "à des versions précises dans requirements.txt."),
        ("A09 — Pas assez de logs",
         "Django enregistre automatiquement les erreurs dans des fichiers de logs. "
         "Sur PythonAnywhere, je peux consulter tous les logs d'accès et d'erreurs. "
         "Les réservations et paiements sont tracés en base de données avec la date et l'heure."),
        ("A10 — Requêtes serveur falsifiées (SSRF)",
         "Sportify ne fait des appels externes que vers l'API Stripe, dont l'adresse est "
         "fixe dans le code. L'application ne fait jamais de requêtes vers des adresses "
         "fournies par les utilisateurs."),
    ]

    for titre, desc in owasp:
        set_heading(doc, titre, 2)
        add_para(doc, desc)

    set_heading(doc, "4. Protection des données personnelles")
    add_para(doc,
        "Sportify collecte uniquement les données nécessaires : nom, prénom, email et "
        "téléphone. Les numéros de carte bancaire ne sont jamais stockés sur Sportify, "
        "ils sont gérés directement par Stripe qui est certifié pour ça. "
        "Les utilisateurs peuvent demander la suppression de leur compte et de leurs données "
        "à tout moment, conformément au RGPD.")

    set_heading(doc, "5. Détection des problèmes (IDS)")
    add_para(doc,
        "Je n'ai pas mis en place un système IDS professionnel car l'hébergement "
        "PythonAnywhere gratuit ne le permet pas. Cependant, plusieurs mécanismes "
        "permettent de détecter des problèmes :")
    add_bullet(doc, "PythonAnywhere enregistre toutes les requêtes HTTP dans des logs accessibles")
    add_bullet(doc, "Django enregistre toutes les erreurs (500, 403, 404) dans un fichier de log")
    add_bullet(doc, "En production réelle, un outil comme Fail2Ban pourrait bloquer automatiquement les adresses IP qui font trop de tentatives de connexion")

    set_heading(doc, "6. Plan de reprise en cas de panne (Disaster Recovery Plan)")
    add_para(doc,
        "Si le site venait à tomber ou à être compromis, voici comment je pourrais le remettre "
        "en marche rapidement :")
    add_bullet(doc, "Le code est sur GitHub : je peux le récupérer et le redéployer en quelques minutes")
    add_bullet(doc, "La base de données peut être sauvegardée régulièrement via un export SQL")
    add_bullet(doc, "Étapes de restauration : identifier le problème, restaurer la dernière sauvegarde, redéployer depuis GitHub, vérifier que tout fonctionne")
    add_para(doc,
        "L'objectif serait de remettre le site en ligne en moins de 4 heures et de ne "
        "perdre au maximum que 24 heures de données.")

    set_heading(doc, "7. Security By Design")
    add_para(doc,
        "Le Security By Design est une approche qui consiste à penser à la sécurité "
        "dès le début du projet, et pas seulement à la fin. Dans Sportify, j'ai appliqué "
        "ce principe de plusieurs façons :")
    add_bullet(doc, "Les rôles utilisateurs ont été définis avant même de créer les pages")
    add_bullet(doc, "Chaque formulaire valide les données avant de les traiter")
    add_bullet(doc, "Par défaut, l'accès est refusé : si une vérification échoue, l'utilisateur ne passe pas")
    add_bullet(doc, "Le code est organisé de façon claire (modèles, vues, templates séparés) pour éviter les erreurs")

    set_heading(doc, "8. Conclusion")
    add_para(doc,
        "Grâce à Django et aux bonnes pratiques de développement, Sportify est protégé "
        "contre les principales menaces web. Le framework gère automatiquement beaucoup "
        "d'aspects de sécurité, ce qui m'a permis de me concentrer sur la logique "
        "de l'application tout en ayant une base solide. "
        "Des améliorations futures pourraient inclure la double authentification (2FA) "
        "et des audits de sécurité réguliers.")

    path = os.path.join(OUTPUT_DIR, 'T16_Strategie_Securite.docx')
    doc.save(path)
    print('[OK] T16_Strategie_Securite.docx')


# ============================================================
#  T17 - STRATEGIE DE REFERENCEMENT
# ============================================================
def create_t17():
    doc = Document()
    title_page(doc, "Stratégie de Référencement", "Travail 17 — Application Web Sportify")

    set_heading(doc, "1. Introduction")
    add_para(doc,
        "Pour que Sportify soit visible sur internet et attire des utilisateurs, "
        "il faut travailler sur le référencement. Le référencement c'est l'ensemble "
        "des techniques qui permettent à un site d'apparaître dans les résultats de recherche "
        "Google ou sur les réseaux sociaux. Il existe 3 types : le SEO (naturel), "
        "le SEA (payant) et le SMO (réseaux sociaux). Dans ce document j'explique "
        "la stratégie que j'adopterais pour Sportify.")

    set_heading(doc, "2. SEO — Le référencement naturel (gratuit)")
    add_para(doc,
        "Le SEO c'est ce qu'on fait pour apparaître dans les résultats Google sans payer. "
        "Ça prend du temps mais les effets durent longtemps.")

    set_heading(doc, "2.1 Ce que j'ai mis en place dans le code", 2)
    add_bullet(doc, "Chaque page a un titre <title> différent et descriptif. Par exemple la page des terrains a le titre \"Terrains disponibles — Sportify\"")
    add_bullet(doc, "Les balises <meta description> résument le contenu de chaque page en 1-2 phrases pour que Google puisse l'afficher dans les résultats")
    add_bullet(doc, "J'utilise les balises HTML sémantiques correctes : <header>, <main>, <footer>, <h1>, <h2>... ce qui aide Google à comprendre la structure du site")
    add_bullet(doc, "Toutes les images ont un attribut alt qui décrit ce qu'elles représentent")
    add_bullet(doc, "Le site est responsive, c'est-à-dire qu'il s'adapte aux téléphones. Google pénalise les sites qui ne le sont pas.")

    set_heading(doc, "2.2 Les URLs", 2)
    add_para(doc, "Les adresses des pages sont courtes et lisibles :")
    add_bullet(doc, "/terrains/ pour voir tous les terrains")
    add_bullet(doc, "/reserver/12/ pour réserver un créneau")
    add_para(doc, "Ce genre d'URL est mieux compris par Google qu'une adresse comme /page?id=3&type=2.")

    set_heading(doc, "2.3 Les mots-clés ciblés", 2)
    add_para(doc, "Pour que Sportify remonte dans Google quand quelqu'un cherche un terrain, "
        "j'ai identifié les mots-clés que les gens tapent :")
    add_bullet(doc, "\"réserver terrain football indoor Belgique\"")
    add_bullet(doc, "\"location terrain padel Bruxelles\"")
    add_bullet(doc, "\"terrain de futsal à Liège pas cher\"")
    add_para(doc, "Ces mots-clés sont intégrés naturellement dans les textes des pages.")

    set_heading(doc, "2.4 Fichiers techniques", 2)
    add_bullet(doc, "Un fichier robots.txt indique à Google quelles pages il peut indexer")
    add_bullet(doc, "Un fichier sitemap.xml liste toutes les pages du site pour que Google les trouve plus facilement")
    add_bullet(doc, "Google Search Console permet de suivre comment le site est indexé et quels mots-clés amènent des visiteurs")

    set_heading(doc, "3. SEA — La publicité payante")
    add_para(doc,
        "Le SEA c'est quand on paie pour apparaître en haut des résultats Google. "
        "C'est plus rapide que le SEO mais ça coûte de l'argent et ça s'arrête "
        "dès qu'on arrête de payer.")

    set_heading(doc, "3.1 Google Ads", 2)
    add_para(doc, "Pour Sportify, je mettrais en place des campagnes Google Ads avec :")
    add_bullet(doc, "Des annonces textuelles qui apparaissent quand quelqu'un cherche \"réserver terrain foot Bruxelles\"")
    add_bullet(doc, "Un ciblage géographique limité à la Belgique pour ne pas gaspiller le budget")
    add_bullet(doc, "Un budget de départ d'environ 10-20€ par jour pour tester ce qui fonctionne")

    set_heading(doc, "3.2 Retargeting", 2)
    add_para(doc,
        "Le retargeting permet de ré-afficher des publicités aux personnes qui ont visité "
        "Sportify sans réserver. Par exemple si quelqu'un a regardé un terrain mais n'a "
        "pas réservé, il verrait une pub Sportify sur d'autres sites pour lui rappeler.")

    set_heading(doc, "3.3 Ce qu'on mesure", 2)
    add_bullet(doc, "Le taux de clics sur les annonces (CTR)")
    add_bullet(doc, "Le coût par clic (CPC) pour savoir si c'est rentable")
    add_bullet(doc, "Le nombre de réservations générées par les publicités")

    set_heading(doc, "4. SMO — Les réseaux sociaux")
    add_para(doc,
        "Le SMO c'est tout ce qui concerne la présence de Sportify sur les réseaux sociaux. "
        "L'objectif est d'attirer des utilisateurs via Instagram, Facebook ou TikTok.")

    set_heading(doc, "4.1 Présence sur les réseaux", 2)
    add_bullet(doc, "Instagram : photos des terrains, vidéos de matchs, témoignages de membres")
    add_bullet(doc, "Facebook : page officielle avec les horaires, les prix et les événements")
    add_bullet(doc, "TikTok : courtes vidéos dynamiques pour toucher les jeunes")

    set_heading(doc, "4.2 Stratégie de contenu", 2)
    add_bullet(doc, "3 à 4 publications par semaine")
    add_bullet(doc, "Hashtags : #SportifyBelgique #FootballIndoor #PadelBruxelles #ReservationTerrain")
    add_bullet(doc, "Encourager les membres à partager leurs matchs avec #MonMatchSportify")

    set_heading(doc, "4.3 Balises Open Graph", 2)
    add_para(doc,
        "Quand quelqu'un partage une page de Sportify sur Facebook ou WhatsApp, "
        "les balises Open Graph permettent d'afficher automatiquement une belle "
        "image, un titre et une description. J'ai intégré ces balises dans le code HTML.")

    set_heading(doc, "5. Outils de suivi")
    add_bullet(doc, "Google Analytics : pour voir combien de personnes visitent le site et d'où elles viennent")
    add_bullet(doc, "Google Search Console : pour voir quels mots-clés amènent des visiteurs")
    add_bullet(doc, "Meta Business Suite : pour gérer et analyser les publicités Facebook/Instagram")

    set_heading(doc, "6. Conclusion")
    add_para(doc,
        "Le référencement de Sportify repose sur les 3 piliers : le SEO pour construire "
        "une visibilité durable sur le long terme, le SEA pour attirer rapidement des "
        "utilisateurs au lancement, et le SMO pour créer une communauté autour de la marque. "
        "Ces 3 approches sont complémentaires et permettent de toucher différents types "
        "d'utilisateurs.")

    path = os.path.join(OUTPUT_DIR, 'T17_Strategie_Referencement.docx')
    doc.save(path)
    print('[OK] T17_Strategie_Referencement.docx')


# ============================================================
#  T18 - ASPECTS JURIDIQUES
# ============================================================
def create_t18():
    doc = Document()
    title_page(doc, "Aspects Juridiques et Cadre Légal", "Travail 18 — Application Web Sportify")

    set_heading(doc, "1. Introduction")
    add_para(doc,
        "Sportify est un site web qui fonctionne en Belgique et qui gère des comptes "
        "utilisateurs et des paiements en ligne. Cela signifie qu'il doit respecter "
        "plusieurs lois : le RGPD pour la protection des données personnelles, "
        "les règles du commerce électronique belge, et le droit d'auteur. "
        "Dans ce document je présente les obligations légales de Sportify et "
        "comment je les respecte.")

    set_heading(doc, "2. Le RGPD — Protection des données personnelles")
    add_para(doc,
        "Le RGPD (Règlement Général sur la Protection des Données) est une loi européenne "
        "qui oblige les sites web à protéger les données personnelles de leurs utilisateurs. "
        "Elle est en vigueur depuis 2018 dans toute l'Union Européenne, donc aussi en Belgique.")

    set_heading(doc, "2.1 Quelles données Sportify collecte", 2)
    add_para(doc, "Quand un utilisateur crée un compte sur Sportify, je collecte :")
    add_bullet(doc, "Son nom et prénom")
    add_bullet(doc, "Son adresse email")
    add_bullet(doc, "Son numéro de téléphone")
    add_bullet(doc, "Son mot de passe (stocké de façon sécurisée, jamais en clair)")
    add_para(doc,
        "Je ne collecte pas plus d'informations que nécessaire. "
        "Les données bancaires sont gérées directement par Stripe, "
        "je ne les stocke pas sur Sportify.")

    set_heading(doc, "2.2 Les droits des utilisateurs", 2)
    add_para(doc, "Grâce au RGPD, chaque utilisateur de Sportify a le droit de :")
    add_bullet(doc, "Consulter ses données personnelles stockées")
    add_bullet(doc, "Corriger des informations incorrectes")
    add_bullet(doc, "Demander la suppression de son compte et de ses données")
    add_bullet(doc, "S'opposer à l'utilisation de ses données à des fins publicitaires")
    add_para(doc,
        "Ces demandes peuvent être faites via la page profil ou par email. "
        "Je m'engage à répondre dans un délai de 30 jours.")

    set_heading(doc, "2.3 Combien de temps les données sont conservées", 2)
    add_bullet(doc, "Les données de compte : tant que l'utilisateur a un compte actif")
    add_bullet(doc, "Les données de paiement : 10 ans (obligation légale en Belgique)")
    add_bullet(doc, "Les logs de connexion : 12 mois maximum")

    set_heading(doc, "3. Les cookies")
    add_para(doc,
        "Les cookies sont de petits fichiers stockés dans le navigateur de l'utilisateur. "
        "Sportify en utilise plusieurs types :")
    add_bullet(doc, "Cookies essentiels : nécessaires au fonctionnement du site (connexion, sécurité). Pas besoin de consentement pour ces cookies.")
    add_bullet(doc, "Cookies analytiques : Google Analytics pour compter les visiteurs. L'utilisateur doit accepter avant qu'ils soient activés.")
    add_bullet(doc, "Cookies publicitaires : pour le retargeting. L'utilisateur doit aussi accepter.")
    add_para(doc,
        "Une bannière de consentement s'affiche lors de la première visite pour informer "
        "l'utilisateur et lui permettre d'accepter ou de refuser les cookies non essentiels.")

    set_heading(doc, "4. Mentions légales")
    add_para(doc, "Conformément à la loi belge, Sportify affiche les informations suivantes :")
    add_bullet(doc, "Responsable : EL HALIMI Morad")
    add_bullet(doc, "Statut : projet académique, étudiant en Bac Info Appli")
    add_bullet(doc, "Email de contact : contact@sportify.be")
    add_bullet(doc, "Hébergeur : PythonAnywhere Ltd, Londres, Royaume-Uni")
    add_bullet(doc, "URL : https://ehsportify.pythonanywhere.com")

    set_heading(doc, "5. Conditions Générales d'Utilisation (CGU)")
    add_para(doc,
        "Les CGU sont les règles que l'utilisateur accepte quand il crée son compte. "
        "Voici les points principaux :")

    set_heading(doc, "5.1 Qui peut s'inscrire", 2)
    add_bullet(doc, "Toute personne majeure (18 ans minimum)")
    add_bullet(doc, "Les informations fournies doivent être exactes et réelles")
    add_bullet(doc, "Un seul compte par personne, le compte n'est pas transférable")

    set_heading(doc, "5.2 Réservations et annulations", 2)
    add_bullet(doc, "Une réservation confirmée et payée est définitive")
    add_bullet(doc, "Annulation possible plus de 24h avant le créneau : remboursement complet")
    add_bullet(doc, "Annulation moins de 24h avant : pas de remboursement")
    add_bullet(doc, "Si c'est Sportify qui annule (cas exceptionnel) : remboursement complet garanti")

    set_heading(doc, "5.3 Ce qui est interdit", 2)
    add_bullet(doc, "Utiliser le site pour des activités illégales")
    add_bullet(doc, "Essayer de pirater ou d'accéder à des données qui ne vous appartiennent pas")
    add_bullet(doc, "Créer de faux comptes ou fausses réservations")
    add_para(doc, "Tout abus entraîne la suppression immédiate du compte.")

    set_heading(doc, "6. Les paiements en ligne")
    add_para(doc,
        "Les paiements sur Sportify sont traités par Stripe, "
        "un service de paiement sécurisé utilisé par des millions de sites dans le monde. "
        "Stripe est certifié PCI-DSS (norme de sécurité pour les paiements en ligne). "
        "Sportify ne voit jamais et ne stocke jamais les numéros de carte bancaire.")
    add_bullet(doc, "Paiements acceptés : carte bancaire Visa et Mastercard")
    add_bullet(doc, "Devise : Euro (EUR)")
    add_bullet(doc, "Une confirmation de paiement est envoyée par email après chaque transaction")
    add_bullet(doc, "Droit de rétractation de 14 jours sauf si le créneau a déjà eu lieu")

    set_heading(doc, "7. Responsabilité")
    add_para(doc,
        "Sportify est une plateforme intermédiaire entre les membres et les propriétaires "
        "de terrains. Cela signifie que :")
    add_bullet(doc, "Sportify n'est pas responsable des accidents ou dommages qui surviendraient sur un terrain")
    add_bullet(doc, "Sportify essaie de garantir la disponibilité du site mais ne peut pas promettre un fonctionnement 24h/24 sans aucune interruption")
    add_bullet(doc, "En cas de force majeure (panne générale, catastrophe naturelle) Sportify ne peut pas être tenu responsable")

    set_heading(doc, "8. Droit d'auteur")
    add_para(doc,
        "Tout ce qui compose Sportify (le code, le design, les textes, le logo) "
        "est protégé par le droit d'auteur belge.")
    add_bullet(doc, "Le code source appartient à EL HALIMI Morad")
    add_bullet(doc, "Le logo et la charte graphique sont des créations originales")
    add_bullet(doc, "Les images utilisées viennent de sites libres de droits (Unsplash, Pexels)")
    add_bullet(doc, "La police de caractères utilisée (Inter) est libre d'utilisation (licence OFL)")
    add_bullet(doc, "Il est interdit de copier ou d'utiliser le contenu de Sportify sans autorisation")

    set_heading(doc, "9. Quel droit s'applique")
    add_para(doc,
        "Les CGU de Sportify sont soumises au droit belge. "
        "En cas de litige qui ne peut pas être réglé à l'amiable, "
        "les tribunaux de Bruxelles seraient compétents. "
        "Les utilisateurs peuvent aussi utiliser la plateforme européenne de résolution "
        "de litiges en ligne (ODR) pour les conflits de consommation.")

    set_heading(doc, "10. Conclusion")
    add_para(doc,
        "Sportify respecte les principales obligations légales qui s'appliquent à un site "
        "de commerce électronique en Belgique. Le RGPD est respecté dans la collecte et "
        "le traitement des données, les paiements sont sécurisés via Stripe, "
        "et les CGU définissent clairement les droits et obligations de chacun. "
        "Ce document sera mis à jour si la législation évolue.")

    path = os.path.join(OUTPUT_DIR, 'T18_Aspects_Juridiques.docx')
    doc.save(path)
    print('[OK] T18_Aspects_Juridiques.docx')


if __name__ == '__main__':
    create_t16()
    create_t17()
    create_t18()
    print('\nDone :', OUTPUT_DIR)
